"""
Resume Parser Service
Extracts text from PDF, DOCX, and TXT files
"""
from io import BytesIO
import re


def parse_resume_file(content: bytes, content_type: str) -> str:
    """
    Parse resume file and extract text based on file type
    
    Args:
        content: File content as bytes
        content_type: MIME type of the file
        
    Returns:
        Extracted text from the file
    """
    if content_type == "application/pdf":
        return extract_text_from_pdf(content)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(content)
    elif content_type == "text/plain":
        return content.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file type: {content_type}")


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """
    Extract text from PDF file using PyPDF2
    """
    try:
        from PyPDF2 import PdfReader
        
        pdf_file = BytesIO(pdf_bytes)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(text)
                print(f"Extracted {len(text)} characters from page {page_num + 1}")
        
        full_text = "\n".join(text_parts)
        cleaned_text = clean_text(full_text)
        
        print(f"Total extracted text: {len(cleaned_text)} characters")
        return cleaned_text
        
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF parsing. Install it with: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"Error parsing PDF: {str(e)}")


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """
    Extract text from DOCX file using python-docx
    """
    try:
        from docx import Document
        
        docx_file = BytesIO(docx_bytes)
        doc = Document(docx_file)
        
        text_parts = []
     
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
       
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = "\n".join(text_parts)
        cleaned_text = clean_text(full_text)
        
        print(f"Extracted text from DOCX: {len(cleaned_text)} characters")
        return cleaned_text
        
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install it with: pip install python-docx")
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text
    """

    text = re.sub(r'\s+', ' ', text)
    
    text = re.sub(r'[^\w\s\-.,;:()\[\]@#+=/"\'&%$]', '', text)
    
   
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()